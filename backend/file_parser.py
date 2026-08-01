"""
file_parser.py
~~~~~~~~~~~~~~
Utility module for parsing multi-format file uploads (CSV, PDF, DOCX, images)
and extracting structured records such as name/email pairs or matric numbers.

Exports one main function:  extract_records_from_file()
"""

from __future__ import annotations

import csv
import io
import re
import tempfile
from pathlib import Path
from typing import Any


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

ALLOWED_EXTENSIONS = {".csv", ".pdf", ".docx", ".doc", ".jpg", ".jpeg", ".png", ".webp"}

IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".webp"}

# Common header aliases -> canonical column name
COLUMN_ALIASES: dict[str, str] = {
    "name": "full_name",
    "student_name": "full_name",
    "matric_no": "matric_number",
    "matricno": "matric_number",
    "matric": "matric_number",
    "student_id": "matric_number",
    "mail": "email",
    # Enrollment aliases
    "course_cd": "course_code",
    "coursecode": "course_code",
    "code": "course_code",
    "course_name": "course_title",
    "course_nm": "course_title",
    "title": "course_title",
    "units": "course_unit",
    "credit": "course_unit",
    "credit_unit": "course_unit",
    "credit_units": "course_unit",
    "unit": "course_unit",
}

# Regex helpers
EMAIL_RE = re.compile(r"[\w.-]+@[\w.-]+\.\w+")
# Rough heuristic: two or more capitalised words (e.g. "John Doe")
NAME_RE = re.compile(r"[A-Z][a-z]+(?:\s+[A-Z][a-z]+)+")


# ---------------------------------------------------------------------------
# Public helpers
# ---------------------------------------------------------------------------

def get_file_extension(filename: str) -> str:
    """Return the lowercased extension if it is in the allow-list, else raise."""
    ext = Path(filename).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type: {ext}. "
            f"Accepted: CSV, PDF, DOCX, JPG, PNG, WEBP"
        )
    return ext

# ---------------------------------------------------------------------------
# Main entry point
# ---------------------------------------------------------------------------

def extract_records_from_file(
    file_bytes: bytes,
    filename: str,
    expected_columns: list[str],
) -> dict[str, Any]:
    """
    Parse an uploaded file and extract records.

    Args:
        file_bytes:        Raw file content.
        filename:          Original filename (used for extension detection).
        expected_columns:  List of column names to extract,
                           e.g. ['full_name', 'email'] or
                           ['matric_number', 'full_name'].

    Returns:
        A dict with the following keys::

            {
                "entries":  [{"full_name": "...", "email": "..."}, ...],
                "errors":   ["Row 3: missing email", ...],
                "requires_manual_review": False,
                "raw_text": None,
            }
    """
    ext = get_file_extension(filename)

    if ext == ".csv":
        return _parse_csv(file_bytes, expected_columns)
    elif ext == ".pdf":
        return _parse_pdf(file_bytes, expected_columns)
    elif ext in {".docx", ".doc"}:
        return _parse_docx(file_bytes, expected_columns)
    elif ext in IMAGE_EXTENSIONS:
        return _parse_image()
    else:
        # Should be unreachable because get_file_extension validates.
        raise ValueError(f"Unsupported file type: {ext}")

# ---------------------------------------------------------------------------
# CSV parser
# ---------------------------------------------------------------------------

def _normalize_header(header: str) -> str:
    """Strip whitespace, lowercase, normalize separators, then resolve aliases."""
    key = header.strip().lower().replace(" ", "_").replace("-", "_")
    return COLUMN_ALIASES.get(key, key)


def _parse_csv(file_bytes: bytes, expected_columns: list[str]) -> dict[str, Any]:
    entries: list[dict[str, str]] = []
    errors: list[str] = []

    text = file_bytes.decode("utf-8-sig")  # utf-8-sig handles optional BOM
    reader = csv.DictReader(io.StringIO(text))

    if reader.fieldnames is None:
        return {
            "entries": [],
            "errors": ["CSV file appears to be empty or has no header row."],
            "requires_manual_review": False,
            "raw_text": None,
        }

    # Build a mapping from the normalised header to the original header
    header_map: dict[str, str] = {}
    for original in reader.fieldnames:
        normalised = _normalize_header(original)
        header_map[normalised] = original

    for row_idx, row in enumerate(reader, start=2):  # row 1 is the header
        record: dict[str, str] = {}
        row_errors: list[str] = []

        for col in expected_columns:
            original_key = header_map.get(col)
            if original_key is None:
                # Column not present in CSV at all - skip silently (reported once below)
                continue
            value = (row.get(original_key) or "").strip()
            if not value:
                row_errors.append(f"Row {row_idx}: missing {col}")
            else:
                record[col] = value

        if row_errors:
            errors.extend(row_errors)
            continue  # skip rows with missing required fields

        # Only add the record if we got at least one expected column
        if record:
            entries.append(record)

    # Warn once if an expected column was not found at all
    for col in expected_columns:
        if col not in header_map:
            errors.insert(0, f"Column \'{col}\' not found in CSV headers.")

    return {
        "entries": entries,
        "errors": errors,
        "requires_manual_review": False,
        "raw_text": None,
    }

# ---------------------------------------------------------------------------
# PDF parser
# ---------------------------------------------------------------------------

def _parse_pdf(file_bytes: bytes, expected_columns: list[str]) -> dict[str, Any]:
    try:
        import pdfplumber  # type: ignore
    except ImportError:
        return {
            "entries": [],
            "errors": [
                "PDF parsing requires the \'pdfplumber\' package. "
                "Install it with: pip install pdfplumber"
            ],
            "requires_manual_review": False,
            "raw_text": None,
        }

    entries: list[dict[str, str]] = []
    errors: list[str] = []
    all_text_parts: list[str] = []
    tmp_path: str | None = None

    try:
        # Write bytes to a temp file because pdfplumber needs a file path
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        with pdfplumber.open(tmp_path) as pdf:
            table_found = False

            for page in pdf.pages:
                # --- attempt tabular extraction first ---
                tables = page.extract_tables()
                if tables:
                    table_found = True
                    for table in tables:
                        if not table:
                            continue
                        # First row is assumed to be header
                        raw_headers = [
                            _normalize_header(h) if h else ""
                            for h in table[0]
                        ]
                        for data_row_idx, data_row in enumerate(
                            table[1:], start=2
                        ):
                            record: dict[str, str] = {}
                            row_errors: list[str] = []
                            for col in expected_columns:
                                if col in raw_headers:
                                    idx = raw_headers.index(col)
                                    value = (
                                        (data_row[idx] or "").strip()
                                        if idx < len(data_row)
                                        else ""
                                    )
                                    if not value:
                                        row_errors.append(
                                            f"Row {data_row_idx}: missing {col}"
                                        )
                                    else:
                                        record[col] = value
                            if row_errors:
                                errors.extend(row_errors)
                                continue
                            if record:
                                entries.append(record)

                # --- always collect page text for fallback ---
                page_text = page.extract_text()
                if page_text:
                    all_text_parts.append(page_text)

            full_text = "\n".join(all_text_parts)

            # If no tables were found, try regex-based extraction
            if not table_found:
                entries, errors = _extract_from_text(
                    full_text, expected_columns
                )

            # If we still have no entries, surface the raw text for review
            raw_text: str | None = None
            if not entries and full_text.strip():
                raw_text = full_text

    finally:
        if tmp_path:
            try:
                Path(tmp_path).unlink(missing_ok=True)
            except OSError:
                pass

    return {
        "entries": entries,
        "errors": errors,
        "requires_manual_review": False,
        "raw_text": raw_text if not entries else None,
    }

# ---------------------------------------------------------------------------
# DOCX parser
# ---------------------------------------------------------------------------

def _parse_docx(file_bytes: bytes, expected_columns: list[str]) -> dict[str, Any]:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return {
            "entries": [],
            "errors": [
                "DOCX parsing requires the \'python-docx\' package. "
                "Install it with: pip install python-docx"
            ],
            "requires_manual_review": False,
            "raw_text": None,
        }

    entries: list[dict[str, str]] = []
    errors: list[str] = []
    tmp_path: str | None = None

    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            tmp.write(file_bytes)
            tmp_path = tmp.name

        doc = Document(tmp_path)

        # --- tables ---
        for table in doc.tables:
            if not table.rows:
                continue
            raw_headers = [
                _normalize_header(cell.text) for cell in table.rows[0].cells
            ]
            for data_row_idx, row in enumerate(table.rows[1:], start=2):
                record: dict[str, str] = {}
                row_errors: list[str] = []
                cells = [cell.text.strip() for cell in row.cells]
                for col in expected_columns:
                    if col in raw_headers:
                        idx = raw_headers.index(col)
                        value = cells[idx] if idx < len(cells) else ""
                        if not value:
                            row_errors.append(
                                f"Row {data_row_idx}: missing {col}"
                            )
                        else:
                            record[col] = value
                if row_errors:
                    errors.extend(row_errors)
                    continue
                if record:
                    entries.append(record)

        # --- paragraphs (fallback if no table entries found) ---
        if not entries:
            paragraph_text = "\n".join(
                para.text for para in doc.paragraphs if para.text.strip()
            )
            if paragraph_text.strip():
                entries, para_errors = _extract_from_text(
                    paragraph_text, expected_columns
                )
                errors.extend(para_errors)

    finally:
        if tmp_path:
            try:
                Path(tmp_path).unlink(missing_ok=True)
            except OSError:
                pass

    return {
        "entries": entries,
        "errors": errors,
        "requires_manual_review": False,
        "raw_text": None,
    }

# ---------------------------------------------------------------------------
# Image handler (no OCR)
# ---------------------------------------------------------------------------

def _parse_image() -> dict[str, Any]:
    return {
        "entries": [],
        "errors": [
            "Image uploaded. Please review and enter records manually."
        ],
        "requires_manual_review": True,
        "raw_text": None,
    }


# ---------------------------------------------------------------------------
# Shared regex-based text extraction
# ---------------------------------------------------------------------------

def _extract_from_text(
    text: str, expected_columns: list[str]
) -> tuple[list[dict[str, str]], list[str]]:
    """
    Best-effort extraction of records from free-form text using regex.

    Returns (entries, errors).
    """
    entries: list[dict[str, str]] = []
    errors: list[str] = []

    emails = EMAIL_RE.findall(text)
    names = NAME_RE.findall(text)

    need_email = "email" in expected_columns
    need_name = "full_name" in expected_columns

    if need_email and need_name:
        # Pair names and emails by order of appearance
        count = min(len(names), len(emails))
        for i in range(count):
            entries.append({"full_name": names[i], "email": emails[i]})
        if len(names) != len(emails):
            errors.append(
                f"Found {len(names)} name(s) and {len(emails)} email(s); "
                "some records may be unpaired."
            )
    elif need_email:
        for email in emails:
            entries.append({"email": email})
    elif need_name:
        for name in names:
            entries.append({"full_name": name})

    # Matric numbers (digits, possibly with slashes or hyphens)
    if "matric_number" in expected_columns:
        matric_re = re.compile(r"\b\d{2,4}[/-]\d{2,4}[/-]\d{2,6}\b")
        matrics = matric_re.findall(text)
        if matrics and not entries:
            # No names/emails found; just list matric numbers
            for m in matrics:
                entries.append({"matric_number": m})
        elif matrics:
            # Try to attach matric numbers to existing entries
            for i, m in enumerate(matrics):
                if i < len(entries):
                    entries[i]["matric_number"] = m

    return entries, errors
