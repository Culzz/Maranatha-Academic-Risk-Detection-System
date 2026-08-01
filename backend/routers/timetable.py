"""Timetable upload, parsing, and retrieval router (class, exam, calendar)."""

import os
import re
import uuid
from datetime import datetime, timezone, date as date_type
from fastapi import APIRouter, Depends, HTTPException, File, UploadFile, Form
from typing import Optional
from sqlalchemy.orm import Session

from security import require_role, get_current_user
from database import get_db
from session_utils import get_active_or_latest_session
from realtime import push_event_to_many
import app_models as models

router = APIRouter()

UPLOAD_DIR = os.path.join(os.path.dirname(__file__), "..", "uploads", "timetables")
os.makedirs(UPLOAD_DIR, exist_ok=True)


# ── Venue normalisation ──────────────────────────────────────────────────────

VENUE_MAP = {
    "aud.": "Auditorium", "aud": "Auditorium", "AUD.": "Auditorium", "AUD": "Auditorium",
    "phy lab": "Physics Lab", "PHY Lab": "Physics Lab", "PHY LAB": "Physics Lab",
    "chem lab": "Chemistry Lab", "CHEM LAB": "Chemistry Lab", "Chem Lab": "Chemistry Lab",
    "sci": "Science Block", "SCI": "Science Block",
}


def normalise_venue(raw: str) -> str:
    if not raw:
        return raw
    raw = raw.strip()
    low = raw.lower()

    # Check direct map
    for key, val in VENUE_MAP.items():
        if low == key.lower():
            return val

    # Normalise Lect patterns: "Lect 7", "Lect.H 7", "Lect H7", "Lect.H.7" → "Lect H 7"
    m = re.match(r'Lect\.?\s*H?\.?\s*(\d+)', raw, re.IGNORECASE)
    if m:
        return f"Lect H {m.group(1)}"

    return raw


# ── Helpers ──────────────────────────────────────────────────────────────────

def match_lecturer(db: Session, surname: str):
    """Match a lecturer surname to a user in DB."""
    if not surname:
        return None
    surname = surname.strip()
    # Strip title prefixes
    for prefix in ("Dr.", "Dr", "Mr.", "Mr", "Mrs.", "Mrs", "Engr.", "Engr", "Prof.", "Prof", "Miss", "Miss."):
        if surname.lower().startswith(prefix.lower()):
            surname = surname[len(prefix):].strip()

    if not surname:
        return None

    lecturers = db.query(models.User).filter(models.User.role == "lecturer").all()
    for lec in lecturers:
        last_name = lec.full_name.split()[-1] if lec.full_name else ""
        if last_name.lower() == surname.lower():
            return lec
    return None


def match_course(db: Session, code_raw: str, session_id: int = None):
    """Match a course code to a course in DB."""
    if not code_raw:
        return None
    normalized = re.sub(r'\s+', '', code_raw).upper()

    q = db.query(models.Course)
    if session_id:
        q = q.filter(models.Course.session_id == session_id)

    # Try exact match first
    course = q.filter(
        models.Course.course_code == normalized
    ).first()
    if course:
        return course

    # Try with space: "CSC221" matches "CSC 221"
    course = q.filter(
        models.Course.course_code == re.sub(r'([A-Z]+)', r'\1 ', normalized).strip()
    ).first()
    return course


def get_all_student_and_lecturer_ids(db: Session):
    """Get all student and lecturer user IDs for notifications."""
    users = db.query(models.User.id).filter(
        models.User.role.in_(["student", "lecturer"]),
        models.User.is_active == True,
    ).all()
    return [str(u.id) for u in users]


TIME_SLOTS = ["8am-10am", "10am-12pm", "12pm-1pm", "1pm-3pm", "3pm-5pm"]
DAY_MAP = {
    "MON": "MON", "MONDAY": "MON",
    "TUE": "TUE", "TUESDAY": "TUE", "TUES": "TUE",
    "WED": "WED", "WEDNESDAY": "WED",
    "THURS": "THURS", "THURSDAY": "THURS", "THU": "THURS",
    "FRI": "FRI", "FRIDAY": "FRI",
}


# ══════════════════════════════════════════════════════════════════════════════
# CLASS TIMETABLE
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/class/upload", status_code=201)
def upload_class_timetable(
    file: UploadFile = File(...),
    session_id: int = Form(...),
    current_user: models.User = Depends(require_role("admin")),
    db: Session = Depends(get_db),
):
    """Upload a DOCX class timetable. Parse tables and insert entries."""
    if not file.filename.endswith((".docx", ".doc")):
        raise HTTPException(400, "Only DOCX files are supported for class timetable.")

    from docx import Document

    saved_name = f"class_{session_id}_{uuid.uuid4().hex[:8]}_{file.filename}"
    saved_path = os.path.join(UPLOAD_DIR, saved_name)
    with open(saved_path, "wb") as f:
        f.write(file.file.read())

    doc = Document(saved_path)

    # Deactivate previous timetable entries for this session
    db.query(models.ClassTimetable).filter(
        models.ClassTimetable.session_id == session_id,
        models.ClassTimetable.is_active == True,
    ).update({"is_active": False})

    inserted = 0
    unmatched_courses = set()
    unmatched_lecturers = set()

    # Try to extract department/faculty from paragraphs before tables
    current_dept = None
    current_faculty = None

    # Build dynamic lookup from DB instead of hardcoding department names
    all_depts = db.query(models.Department).all()
    all_faculties = db.query(models.Faculty).all()

    para_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for pt in para_texts:
        upper = pt.upper()
        # Match department names dynamically
        for dept in all_depts:
            if dept.name.upper() in upper or dept.code.upper() in upper:
                current_dept = dept.name
                break
        # Match faculty names/codes dynamically
        for fac in all_faculties:
            if fac.code.upper() in upper or fac.name.upper() in upper:
                current_faculty = fac.code
                break

    # Parse each table
    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue

        # First row is header — detect time slots
        header_cells = [cell.text.strip() for cell in rows[0].cells]
        slot_map = {}  # col_index -> time_slot string
        for col_idx, cell_text in enumerate(header_cells):
            if col_idx == 0:
                continue  # DAYS column
            # Normalise time slot text
            slot = cell_text.strip()
            if slot:
                slot_map[col_idx] = slot

        # Process data rows
        for row in rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if not cells:
                continue

            # First cell = day
            day_text = cells[0].upper().strip()
            day = DAY_MAP.get(day_text, day_text)
            if day not in ("MON", "TUE", "WED", "THURS", "FRI"):
                continue

            for col_idx, cell_text in enumerate(cells[1:], start=1):
                if not cell_text:
                    continue
                time_slot = slot_map.get(col_idx, f"slot_{col_idx}")

                # Check for BREAK
                if cell_text.upper().strip() == "BREAK":
                    db.add(models.ClassTimetable(
                        session_id=session_id,
                        department=current_dept,
                        faculty=current_faculty,
                        day_of_week=day,
                        time_slot=time_slot,
                        course_code="BREAK",
                        is_break=True,
                        uploaded_by=current_user.id,
                    ))
                    inserted += 1
                    continue

                # Parse entries: "CSC 111(Iluebe, Lect 7)" or "MTH 111 (Iluebe, Lect 7)"
                # Multiple entries may be separated by newlines
                entries = cell_text.split('\n')
                for entry in entries:
                    entry = entry.strip()
                    if not entry:
                        continue

                    # Try pattern: COURSE_CODE (Lecturer, Venue)
                    pattern = r'([A-Z]{2,4}\s?\d{3}[A-Z]?)\s*\(([^,)]+)(?:,\s*([^)]+))?\)'
                    matches = re.findall(pattern, entry, re.IGNORECASE)

                    if matches:
                        for m in matches:
                            course_code = m[0].strip()
                            lecturer_or_venue = m[1].strip()
                            venue_raw = m[2].strip() if m[2] else None

                            # If no venue, the second group might be venue, not lecturer
                            lecturer_name = lecturer_or_venue
                            venue = normalise_venue(venue_raw) if venue_raw else None

                            # Match course
                            course = match_course(db, course_code, session_id)
                            course_id = course.id if course else None
                            if not course:
                                unmatched_courses.add(course_code)

                            # Match lecturer
                            lec = match_lecturer(db, lecturer_name)
                            lecturer_id = lec.id if lec else None
                            if not lec and lecturer_name:
                                unmatched_lecturers.add(lecturer_name)

                            # If course matched, get title for reference
                            db.add(models.ClassTimetable(
                                session_id=session_id,
                                department=current_dept,
                                faculty=current_faculty,
                                day_of_week=day,
                                time_slot=time_slot,
                                course_id=course_id,
                                course_code=course_code,
                                lecturer_id=lecturer_id,
                                lecturer_name_raw=lecturer_name,
                                venue=venue,
                                is_break=False,
                                uploaded_by=current_user.id,
                            ))
                            inserted += 1
                    else:
                        # Try to parse as just a course code
                        code_match = re.match(r'([A-Z]{2,4}\s?\d{3}[A-Z]?)', entry, re.IGNORECASE)
                        if code_match:
                            cc = code_match.group(1).strip()
                            course = match_course(db, cc, session_id)
                            db.add(models.ClassTimetable(
                                session_id=session_id,
                                department=current_dept,
                                faculty=current_faculty,
                                day_of_week=day,
                                time_slot=time_slot,
                                course_id=course.id if course else None,
                                course_code=cc,
                                is_break=False,
                                uploaded_by=current_user.id,
                            ))
                            inserted += 1
                            if not course:
                                unmatched_courses.add(cc)

    # Notify all students and lecturers
    user_ids = get_all_student_and_lecturer_ids(db)
    push_event_to_many(
        db, user_ids, "class_timetable_published",
        {"message": "Your class timetable has been published! Check your Timetable page."}
    )

    db.commit()

    return {
        "inserted": inserted,
        "unmatched_courses": list(unmatched_courses),
        "unmatched_lecturers": list(unmatched_lecturers),
        "message": f"{inserted} timetable entries created.",
    }


@router.patch("/class/{entry_id}")
def update_class_entry(
    entry_id: int,
    payload: dict,
    current_user: models.User = Depends(require_role("admin")),
    db: Session = Depends(get_db),
):
    entry = db.query(models.ClassTimetable).filter(models.ClassTimetable.id == entry_id).first()
    if not entry:
        raise HTTPException(404, "Timetable entry not found.")

    for field in ("course_code", "venue", "lecturer_name_raw", "time_slot", "day_of_week"):
        if field in payload:
            setattr(entry, field, payload[field])

    if "venue" in payload:
        entry.venue = normalise_venue(payload["venue"])

    if "lecturer_name_raw" in payload:
        lec = match_lecturer(db, payload["lecturer_name_raw"])
        entry.lecturer_id = lec.id if lec else None

    if "course_code" in payload:
        course = match_course(db, payload["course_code"], entry.session_id)
        entry.course_id = course.id if course else None

    user_ids = get_all_student_and_lecturer_ids(db)
    push_event_to_many(
        db, user_ids, "class_timetable_updated",
        {"message": "Your class timetable has been updated."}
    )

    db.commit()
    return {"message": "Entry updated."}


@router.delete("/class/{entry_id}")
def delete_class_entry(
    entry_id: int,
    current_user: models.User = Depends(require_role("admin")),
    db: Session = Depends(get_db),
):
    entry = db.query(models.ClassTimetable).filter(models.ClassTimetable.id == entry_id).first()
    if not entry:
        raise HTTPException(404, "Entry not found.")
    db.delete(entry)
    db.commit()
    return {"message": "Entry deleted."}


@router.get("/class/my")
def get_my_class_timetable(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Return class timetable entries for the current user."""
    # Get active session
    active_session = get_active_or_latest_session(db)
    if not active_session:
        return []

    if current_user.role == "lecturer":
        entries = db.query(models.ClassTimetable).filter(
            models.ClassTimetable.session_id == active_session.id,
            models.ClassTimetable.lecturer_id == current_user.id,
            models.ClassTimetable.is_active == True,
        ).order_by(models.ClassTimetable.day_of_week, models.ClassTimetable.time_slot).all()
    else:
        # Student: get enrolled course IDs
        enrolled_ids = [e.course_id for e in db.query(models.Enrollment.course_id).filter(
            models.Enrollment.student_id == current_user.id,
            models.Enrollment.session_id == active_session.id,
        ).all()]

        entries = db.query(models.ClassTimetable).filter(
            models.ClassTimetable.session_id == active_session.id,
            models.ClassTimetable.is_active == True,
            (models.ClassTimetable.course_id.in_(enrolled_ids) | (models.ClassTimetable.is_break == True)),
        ).order_by(models.ClassTimetable.day_of_week, models.ClassTimetable.time_slot).all()

    return [
        {
            "id": e.id,
            "day_of_week": e.day_of_week,
            "time_slot": e.time_slot,
            "course_code": e.course_code,
            "course_title": e.course.course_title if e.course else None,
            "lecturer_name": e.lecturer.full_name if e.lecturer else e.lecturer_name_raw,
            "venue": e.venue,
            "is_break": e.is_break,
            "department": e.department,
        }
        for e in entries
    ]


@router.get("/class/admin")
def get_admin_class_timetable(
    department: Optional[str] = None,
    day: Optional[str] = None,
    current_user: models.User = Depends(require_role("admin")),
    db: Session = Depends(get_db),
):
    active_session = get_active_or_latest_session(db)
    if not active_session:
        return []

    q = db.query(models.ClassTimetable).filter(
        models.ClassTimetable.session_id == active_session.id,
        models.ClassTimetable.is_active == True,
    )
    if department:
        q = q.filter(models.ClassTimetable.department.ilike(f"%{department}%"))
    if day:
        q = q.filter(models.ClassTimetable.day_of_week == day.upper())

    entries = q.order_by(models.ClassTimetable.day_of_week, models.ClassTimetable.time_slot).all()

    return [
        {
            "id": e.id,
            "day_of_week": e.day_of_week,
            "time_slot": e.time_slot,
            "course_code": e.course_code,
            "course_title": e.course.course_title if e.course else None,
            "lecturer_name": e.lecturer.full_name if e.lecturer else e.lecturer_name_raw,
            "venue": e.venue,
            "is_break": e.is_break,
            "department": e.department,
            "faculty": e.faculty,
        }
        for e in entries
    ]


# ══════════════════════════════════════════════════════════════════════════════
# EXAM TIMETABLE
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/exam/upload", status_code=201)
def upload_exam_timetable(
    file: UploadFile = File(...),
    session_id: int = Form(...),
    current_user: models.User = Depends(require_role("admin")),
    db: Session = Depends(get_db),
):
    """Upload a DOCX exam timetable."""
    if not file.filename.endswith((".docx", ".doc")):
        raise HTTPException(400, "Only DOCX files are supported.")

    from docx import Document

    saved_name = f"exam_{session_id}_{uuid.uuid4().hex[:8]}_{file.filename}"
    saved_path = os.path.join(UPLOAD_DIR, saved_name)
    with open(saved_path, "wb") as f:
        f.write(file.file.read())

    doc = Document(saved_path)

    # Deactivate old entries
    db.query(models.ExamTimetable).filter(
        models.ExamTimetable.session_id == session_id,
        models.ExamTimetable.is_active == True,
    ).update({"is_active": False})

    inserted = 0

    for table in doc.tables:
        rows = table.rows
        if len(rows) < 2:
            continue

        # Parse header to find time slot columns
        header_cells = [cell.text.strip() for cell in rows[0].cells]

        for row in rows[1:]:
            cells = [cell.text.strip() for cell in row.cells]
            if not cells or not cells[0]:
                continue

            # First cell: "DAY_NAME\nDD/MM/YY" or "TUESDAY 17/02/26"
            first_cell = cells[0]
            date_match = re.search(r'(\d{1,2})[/\-](\d{1,2})[/\-](\d{2,4})', first_cell)
            exam_date = None
            if date_match:
                day_num = int(date_match.group(1))
                month_num = int(date_match.group(2))
                year_num = int(date_match.group(3))
                if year_num < 100:
                    year_num += 2000
                try:
                    exam_date = date_type(year_num, month_num, day_num)
                except ValueError:
                    pass

            if not exam_date:
                continue

            # Check if Wednesday for community service
            is_wednesday = "WED" in first_cell.upper()

            # Process remaining cells in groups: course | hall | invigilators
            # Typical structure: slot1_course | hall | invig | slot2_course | hall | invig | break | slot3_course | hall | invig
            col_idx = 1
            current_time_slot_idx = 0
            exam_time_slots = ["9:00-11:00", "11:00-1:00", "1:00-2:00", "2:00-4:00"]

            while col_idx < len(cells):
                cell_text = cells[col_idx].strip()

                if not cell_text:
                    col_idx += 1
                    continue

                upper = cell_text.upper()

                # Check for BREAK
                if upper == "BREAK":
                    time_slot = exam_time_slots[min(current_time_slot_idx, len(exam_time_slots) - 1)]
                    db.add(models.ExamTimetable(
                        session_id=session_id,
                        exam_date=exam_date,
                        time_slot=time_slot,
                        course_code="BREAK",
                        is_break=True,
                        uploaded_by=current_user.id,
                    ))
                    inserted += 1
                    current_time_slot_idx += 1
                    col_idx += 1
                    continue

                # Check for community service/fellowship
                if "COMMUNITY" in upper or "FELLOWSHIP" in upper:
                    time_slot = "1:00-2:00"
                    db.add(models.ExamTimetable(
                        session_id=session_id,
                        exam_date=exam_date,
                        time_slot=time_slot,
                        course_code="COMMUNITY SERVICE/FELLOWSHIP",
                        is_community_service=True,
                        uploaded_by=current_user.id,
                    ))
                    inserted += 1
                    current_time_slot_idx += 1
                    col_idx += 1
                    continue

                # Try to parse course codes from cell
                course_codes = re.findall(r'[A-Z]{2,4}\s?\d{3}[A-Z]?', cell_text, re.IGNORECASE)
                if course_codes:
                    time_slot = exam_time_slots[min(current_time_slot_idx, len(exam_time_slots) - 1)]

                    # Next cells might be hall and invigilators
                    hall = cells[col_idx + 1].strip() if col_idx + 1 < len(cells) else None
                    invig_raw = cells[col_idx + 2].strip() if col_idx + 2 < len(cells) else None

                    # If hall looks like a course code, it's not a hall
                    if hall and re.match(r'[A-Z]{2,4}\s?\d{3}', hall):
                        hall = None
                        invig_raw = None

                    hall = normalise_venue(hall) if hall else None

                    for cc in course_codes:
                        course = match_course(db, cc, session_id)
                        exam_entry = models.ExamTimetable(
                            session_id=session_id,
                            exam_date=exam_date,
                            time_slot=time_slot,
                            course_id=course.id if course else None,
                            course_code=cc.strip(),
                            exam_hall=hall,
                            invigilator_names_raw=invig_raw,
                            is_break=False,
                            is_community_service=False,
                            uploaded_by=current_user.id,
                        )
                        db.add(exam_entry)
                        db.flush()

                        # Match invigilators
                        if invig_raw:
                            for name in re.split(r'[,\n]', invig_raw):
                                name = name.strip()
                                if name:
                                    lec = match_lecturer(db, name)
                                    if lec:
                                        db.add(models.ExamTimetableInvigilator(
                                            exam_timetable_id=exam_entry.id,
                                            user_id=lec.id,
                                        ))

                        inserted += 1

                    # Skip hall and invigilator columns
                    col_idx += 3
                    current_time_slot_idx += 1
                else:
                    col_idx += 1

    # Notify all users
    user_ids = get_all_student_and_lecturer_ids(db)
    push_event_to_many(
        db, user_ids, "exam_timetable_published",
        {"message": "Your exam timetable has been published! Check your Timetable page."}
    )

    db.commit()

    return {
        "inserted": inserted,
        "message": f"{inserted} exam timetable entries created.",
    }


@router.get("/exam/my")
def get_my_exam_timetable(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    active_session = get_active_or_latest_session(db)
    if not active_session:
        return []

    if current_user.role == "lecturer":
        # Get exams where they are invigilators
        invig_ids = [i.exam_timetable_id for i in db.query(
            models.ExamTimetableInvigilator.exam_timetable_id
        ).filter(models.ExamTimetableInvigilator.user_id == current_user.id).all()]

        entries = db.query(models.ExamTimetable).filter(
            models.ExamTimetable.session_id == active_session.id,
            models.ExamTimetable.is_active == True,
            (models.ExamTimetable.id.in_(invig_ids) |
             (models.ExamTimetable.is_break == True) |
             (models.ExamTimetable.is_community_service == True)),
        ).order_by(models.ExamTimetable.exam_date, models.ExamTimetable.time_slot).all()
    else:
        # Student: enrolled courses
        enrolled_ids = [e.course_id for e in db.query(models.Enrollment.course_id).filter(
            models.Enrollment.student_id == current_user.id,
            models.Enrollment.session_id == active_session.id,
        ).all()]

        entries = db.query(models.ExamTimetable).filter(
            models.ExamTimetable.session_id == active_session.id,
            models.ExamTimetable.is_active == True,
            (models.ExamTimetable.course_id.in_(enrolled_ids) |
             (models.ExamTimetable.is_break == True) |
             (models.ExamTimetable.is_community_service == True)),
        ).order_by(models.ExamTimetable.exam_date, models.ExamTimetable.time_slot).all()

    return [
        {
            "id": e.id,
            "exam_date": e.exam_date.isoformat() if e.exam_date else None,
            "time_slot": e.time_slot,
            "course_code": e.course_code,
            "course_title": e.course.course_title if e.course else None,
            "exam_hall": e.exam_hall,
            "invigilator_names": e.invigilator_names_raw,
            "is_break": e.is_break,
            "is_community_service": e.is_community_service,
        }
        for e in entries
    ]


# ══════════════════════════════════════════════════════════════════════════════
# ACADEMIC CALENDAR
# ══════════════════════════════════════════════════════════════════════════════

@router.post("/calendar/upload", status_code=201)
def upload_calendar(
    file: UploadFile = File(None),
    calendar_text: Optional[str] = Form(None),
    session_id: int = Form(...),
    current_user: models.User = Depends(require_role("admin")),
    db: Session = Depends(get_db),
):
    """Upload a calendar as PDF or provide raw text. Parse events."""
    text = ""

    if file and file.filename:
        if file.filename.endswith(".pdf"):
            import pdfplumber
            saved_name = f"cal_{session_id}_{uuid.uuid4().hex[:8]}_{file.filename}"
            saved_path = os.path.join(UPLOAD_DIR, saved_name)
            with open(saved_path, "wb") as f:
                f.write(file.file.read())
            with pdfplumber.open(saved_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        elif file.filename.endswith((".txt", ".docx")):
            if file.filename.endswith(".docx"):
                from docx import Document
                saved_name = f"cal_{session_id}_{uuid.uuid4().hex[:8]}_{file.filename}"
                saved_path = os.path.join(UPLOAD_DIR, saved_name)
                with open(saved_path, "wb") as f:
                    f.write(file.file.read())
                doc = Document(saved_path)
                text = "\n".join(p.text for p in doc.paragraphs)
            else:
                text = file.file.read().decode("utf-8", errors="ignore")
        elif file.filename.lower().endswith((".jpg", ".jpeg", ".png", ".webp")):
            try:
                from PIL import Image
                import pytesseract
            except ImportError:
                raise HTTPException(400, "Image OCR requires 'pytesseract' and 'Pillow'. Install with: pip install pytesseract Pillow")
            try:
                img_bytes = file.file.read()
                saved_name = f"cal_{session_id}_{uuid.uuid4().hex[:8]}_{file.filename}"
                saved_path = os.path.join(UPLOAD_DIR, saved_name)
                with open(saved_path, "wb") as f:
                    f.write(img_bytes)
                import io as _io
                img = Image.open(_io.BytesIO(img_bytes))
                text = pytesseract.image_to_string(img)
                if not text.strip():
                    raise HTTPException(400, "Could not extract text from image. Try a clearer photo or enter events manually.")
            except HTTPException:
                raise
            except Exception as exc:
                raise HTTPException(400, f"Image OCR failed: {str(exc)}")
        else:
            raise HTTPException(400, "Supported formats: PDF, DOCX, TXT, JPG, PNG.")

    if calendar_text:
        text = calendar_text

    if not text.strip():
        raise HTTPException(400, "No calendar text provided.")

    # Delete old calendar events for this session
    db.query(models.AcademicCalendarEvent).filter(
        models.AcademicCalendarEvent.session_id == session_id
    ).delete()

    # Parse events from text
    inserted = 0
    current_semester = None
    session_obj = db.query(models.AcademicSession).filter(
        models.AcademicSession.id == session_id
    ).first()
    session_year = 2025  # Fallback
    if session_obj:
        session_year = session_obj.start_date.year if session_obj.start_date else 2025

    months = {
        "january": 1, "february": 2, "march": 3, "april": 4,
        "may": 5, "june": 6, "july": 7, "august": 8,
        "september": 9, "october": 10, "november": 11, "december": 12,
        "jan": 1, "feb": 2, "mar": 3, "apr": 4,
        "jun": 6, "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
    }

    lines = text.split("\n")
    current_month = None

    for line in lines:
        line = line.strip()
        if not line:
            continue

        upper = line.upper()

        # Detect semester header
        if "FIRST SEMESTER" in upper or "1ST SEMESTER" in upper:
            current_semester = "FIRST"
            continue
        elif "SECOND SEMESTER" in upper or "2ND SEMESTER" in upper:
            current_semester = "SECOND"
            continue

        # Detect month header: "September:", "October:", etc.
        month_match = re.match(r'^(january|february|march|april|may|june|july|august|september|october|november|december)\s*:?\s*$', line, re.IGNORECASE)
        if month_match:
            current_month = months.get(month_match.group(1).lower())
            continue

        # Try to parse event lines: "Saturday 5th — Resumption..." or "Monday 7th, Commencement..."
        # Or "5th-7th — Return of students"
        event_pattern = r'(?:(?:Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday)\s+)?(\d{1,2})(?:st|nd|rd|th)?(?:\s*[-–]\s*(\d{1,2})(?:st|nd|rd|th)?)?\s*[-–,]\s*(.+)'
        ev_match = re.match(event_pattern, line, re.IGNORECASE)

        if ev_match and current_month:
            day_start = int(ev_match.group(1))
            day_end = int(ev_match.group(2)) if ev_match.group(2) else None
            event_label = ev_match.group(3).strip()

            # Determine year based on month and semester
            year = session_year
            if current_semester == "SECOND" and current_month >= 3:
                year = session_year + 1
            elif current_semester == "FIRST" and current_month <= 2:
                year = session_year + 1

            try:
                event_date = date_type(year, current_month, day_start)
                event_date_end = date_type(year, current_month, day_end) if day_end else None
            except ValueError:
                continue

            # Classify event type
            event_type = "other"
            label_lower = event_label.lower()
            if "resumption" in label_lower or "return" in label_lower:
                event_type = "resumption"
            elif "exam" in label_lower:
                event_type = "exam"
            elif "break" in label_lower or "christmas" in label_lower or "vacation" in label_lower:
                event_type = "break"
            elif "lecture" in label_lower or "commencement of lecture" in label_lower:
                event_type = "lectures"
            elif "thanksgiving" in label_lower or "fellowship" in label_lower or "service" in label_lower or "spiritual" in label_lower or "carol" in label_lower:
                event_type = "service"
            elif "registration" in label_lower:
                event_type = "lectures"
            elif "matriculation" in label_lower:
                event_type = "other"
            elif "end of" in label_lower:
                event_type = "exam"

            db.add(models.AcademicCalendarEvent(
                session_id=session_id,
                semester=current_semester,
                event_date=event_date,
                event_date_end=event_date_end,
                event_label=event_label,
                event_type=event_type,
                uploaded_by=current_user.id,
            ))
            inserted += 1
        elif current_month is None:
            # Try inline month: "September: Saturday 5th, Resumption..."
            inline_match = re.match(
                r'(january|february|march|april|may|june|july|august|september|october|november|december)\s*:\s*(?:(?:Monday|Tuesday|Wednesday|Thursday|Friday|Saturday|Sunday)\s+)?(\d{1,2})(?:st|nd|rd|th)?(?:\s*[-–]\s*(\d{1,2})(?:st|nd|rd|th)?)?\s*[-–,]\s*(.+)',
                line, re.IGNORECASE
            )
            if inline_match:
                month_num = months.get(inline_match.group(1).lower())
                day_start = int(inline_match.group(2))
                day_end = int(inline_match.group(3)) if inline_match.group(3) else None
                event_label = inline_match.group(4).strip()

                year = session_year
                if current_semester == "SECOND" and month_num >= 3:
                    year = session_year + 1

                try:
                    event_date = date_type(year, month_num, day_start)
                    event_date_end = date_type(year, month_num, day_end) if day_end else None
                except ValueError:
                    continue

                event_type = "other"
                label_lower = event_label.lower()
                if "resumption" in label_lower or "return" in label_lower:
                    event_type = "resumption"
                elif "exam" in label_lower:
                    event_type = "exam"
                elif "break" in label_lower or "christmas" in label_lower:
                    event_type = "break"
                elif "lecture" in label_lower:
                    event_type = "lectures"
                elif "thanksgiving" in label_lower or "fellowship" in label_lower or "service" in label_lower:
                    event_type = "service"

                db.add(models.AcademicCalendarEvent(
                    session_id=session_id,
                    semester=current_semester,
                    event_date=event_date,
                    event_date_end=event_date_end,
                    event_label=event_label,
                    event_type=event_type,
                    uploaded_by=current_user.id,
                ))
                inserted += 1
                current_month = month_num

    # Notify users
    user_ids = get_all_student_and_lecturer_ids(db)
    push_event_to_many(
        db, user_ids, "calendar_updated",
        {"message": "The academic calendar has been updated. Check your Calendar."}
    )

    # ── Auto-create academic sessions from parsed events ──
    sessions_created = []
    try:
        # Gather resumption and exam events per semester
        pending_events = db.new_changes = True  # force flush
        db.flush()  # Make inserted events queryable

        cal_events = db.query(models.AcademicCalendarEvent).filter(
            models.AcademicCalendarEvent.session_id == session_id
        ).all()

        sem_data = {}  # semester_key -> {"resumptions": [], "exams": []}
        for ev in cal_events:
            sem_key = ev.semester or "FIRST"
            if sem_key not in sem_data:
                sem_data[sem_key] = {"resumptions": [], "exams": [], "label": sem_key}
            ev_type = (ev.event_type or "").lower()
            ev_label = (ev.event_label or "").lower()
            if ev_type == "resumption" or "resumption" in ev_label or "return" in ev_label:
                if ev.event_date:
                    sem_data[sem_key]["resumptions"].append(ev.event_date)
            elif ev_type == "exam" or "exam" in ev_label:
                end_d = ev.event_date_end or ev.event_date
                if end_d:
                    sem_data[sem_key]["exams"].append(end_d)

        # Build session_label from the original session object
        label_base = session_obj.session_label if session_obj else f"{session_year}/{session_year + 1}"
        today_ = date_type.today()

        for sem_key, data in sem_data.items():
            if not data["resumptions"]:
                continue
            start = min(data["resumptions"])
            end = max(data["exams"]) if data["exams"] else None
            if not end:
                # Estimate end as start + 20 weeks
                from datetime import timedelta as _td
                end = start + _td(weeks=20)

            if hasattr(start, "date"):
                start = start.date()
            if hasattr(end, "date"):
                end = end.date()

            sem_num = 2 if sem_key == "SECOND" else 1

            # Check if a session with same label + semester already exists
            existing = db.query(models.AcademicSession).filter(
                models.AcademicSession.session_label == label_base,
                models.AcademicSession.semester == sem_num,
            ).first()

            if existing:
                existing.start_date = start
                existing.end_date = end
                created_id = existing.id
            else:
                new_s = models.AcademicSession(
                    session_label=label_base,
                    semester=sem_num,
                    start_date=start,
                    end_date=end,
                    is_active=False,
                )
                db.add(new_s)
                db.flush()
                created_id = new_s.id

            sessions_created.append({
                "id": created_id,
                "label": label_base,
                "semester": sem_num,
                "start": str(start),
                "end": str(end),
            })

        # Auto-activate the session whose date range includes today
        if sessions_created:
            # Deactivate all first
            db.query(models.AcademicSession).update({"is_active": False})
            # Find which semester is current
            activated = None
            for sc in sessions_created:
                from datetime import date as _date_t
                s_start = _date_t.fromisoformat(sc["start"])
                s_end = _date_t.fromisoformat(sc["end"])
                if s_start <= today_ <= s_end:
                    db.query(models.AcademicSession).filter(
                        models.AcademicSession.id == sc["id"]
                    ).update({"is_active": True})
                    activated = sc
                    break
            # If none contains today, activate the earliest future one or latest past one
            if not activated:
                future = [s for s in sessions_created if _date_t.fromisoformat(s["start"]) > today_]
                if future:
                    target = min(future, key=lambda s: s["start"])
                else:
                    target = max(sessions_created, key=lambda s: s["start"])
                db.query(models.AcademicSession).filter(
                    models.AcademicSession.id == target["id"]
                ).update({"is_active": True})
                activated = target

            # Update calendar events to point to correct session IDs
            for sc in sessions_created:
                sem_name = "SECOND" if sc["semester"] == 2 else "FIRST"
                db.query(models.AcademicCalendarEvent).filter(
                    models.AcademicCalendarEvent.session_id == session_id,
                    models.AcademicCalendarEvent.semester == sem_name,
                ).update({"session_id": sc["id"]})

    except Exception as exc:
        import logging
        logging.getLogger("maranatha").warning("Auto-session creation failed: %s", exc)
        # Don't fail the whole upload for this

    db.commit()

    return {
        "inserted": inserted,
        "message": f"{inserted} calendar events extracted.",
        "sessions_created": sessions_created,
    }


@router.post("/calendar/event", status_code=201)
def add_calendar_event(
    payload: dict,
    current_user: models.User = Depends(require_role("admin")),
    db: Session = Depends(get_db),
):
    """Manually add a single calendar event."""
    active_session = get_active_or_latest_session(db)
    if not active_session:
        raise HTTPException(400, "No active session.")

    event = models.AcademicCalendarEvent(
        session_id=active_session.id,
        semester=payload.get("semester"),
        event_date=payload.get("event_date"),
        event_date_end=payload.get("event_date_end"),
        event_label=payload["event_label"],
        event_type=payload.get("event_type", "other"),
        uploaded_by=current_user.id,
    )
    db.add(event)
    db.commit()
    db.refresh(event)
    return {"id": event.id, "message": "Event added."}


@router.delete("/calendar/{event_id}")
def delete_calendar_event(
    event_id: int,
    current_user: models.User = Depends(require_role("admin")),
    db: Session = Depends(get_db),
):
    event = db.query(models.AcademicCalendarEvent).filter(
        models.AcademicCalendarEvent.id == event_id
    ).first()
    if not event:
        raise HTTPException(404, "Event not found.")
    db.delete(event)
    db.commit()
    return {"message": "Event deleted."}


@router.post("/calendar/import-public-holidays")
def import_public_holidays(
    session_id: int = Form(...),
    year: int = Form(None),
    current_user: models.User = Depends(require_role("admin")),
    db: Session = Depends(get_db),
):
    """
    Import Nigerian public holidays from Nager.Date API.
    Creates AcademicCalendarEvent entries with event_type='public_holiday'.
    Idempotent: skips holidays that already exist by label+date.
    """
    import httpx

    session_obj = db.query(models.AcademicSession).filter(
        models.AcademicSession.id == session_id
    ).first()
    if not session_obj:
        raise HTTPException(404, "Session not found.")

    if not year:
        year = session_obj.start_date.year if session_obj.start_date else 2025

    def _fetch_and_insert(fetch_year: int) -> tuple[int, int]:
        inserted, skipped = 0, 0
        try:
            resp = httpx.get(
                f"https://date.nager.at/api/v3/PublicHolidays/{fetch_year}/NG",
                timeout=10.0,
            )
            resp.raise_for_status()
            holidays = resp.json()
        except Exception as exc:
            raise HTTPException(502, f"Failed to fetch public holidays for {fetch_year}: {str(exc)}")

        for h in holidays:
            h_date = date_type.fromisoformat(h["date"])
            h_label = h.get("localName") or h.get("name", "Public Holiday")

            existing = db.query(models.AcademicCalendarEvent).filter(
                models.AcademicCalendarEvent.session_id == session_id,
                models.AcademicCalendarEvent.event_date == h_date,
                models.AcademicCalendarEvent.event_label == h_label,
            ).first()
            if existing:
                skipped += 1
                continue

            db.add(models.AcademicCalendarEvent(
                session_id=session_id,
                semester=None,
                event_date=h_date,
                event_date_end=None,
                event_label=h_label,
                event_type="public_holiday",
                uploaded_by=current_user.id,
            ))
            inserted += 1
        return inserted, skipped

    total_inserted, total_skipped = _fetch_and_insert(year)

    # If session spans two years, also fetch the next year
    years_fetched = [year]
    if session_obj.end_date and session_obj.end_date.year > year:
        next_year = year + 1
        try:
            ins, sk = _fetch_and_insert(next_year)
            total_inserted += ins
            total_skipped += sk
            years_fetched.append(next_year)
        except Exception:
            pass  # Best-effort for second year

    db.commit()
    year_str = " and ".join(str(y) for y in years_fetched)
    return {
        "inserted": total_inserted,
        "skipped": total_skipped,
        "message": f"Imported {total_inserted} public holidays for {year_str}.",
    }


@router.get("/calendar")
def get_calendar_events(
    semester: Optional[str] = None,
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    active_session = get_active_or_latest_session(db)
    if not active_session:
        return []

    q = db.query(models.AcademicCalendarEvent).filter(
        models.AcademicCalendarEvent.session_id == active_session.id
    )
    if semester:
        q = q.filter(models.AcademicCalendarEvent.semester == semester.upper())

    events = q.order_by(models.AcademicCalendarEvent.event_date).all()

    return [
        {
            "id": e.id,
            "semester": e.semester,
            "event_date": e.event_date.isoformat() if e.event_date else None,
            "event_date_end": e.event_date_end.isoformat() if e.event_date_end else None,
            "event_label": e.event_label,
            "event_type": e.event_type,
        }
        for e in events
    ]


@router.get("/class/my/export.ics")
def export_class_timetable_ics(
    current_user: models.User = Depends(get_current_user),
    db: Session = Depends(get_db),
):
    """Export student's class timetable as .ics (iCalendar) file."""
    from fastapi.responses import Response
    from datetime import timedelta

    if current_user.role == "student":
        enrolled = db.query(models.Enrollment.course_id).filter(
            models.Enrollment.student_id == current_user.id,
        ).all()
        course_ids = [e.course_id for e in enrolled]
        entries = db.query(models.ClassTimetableEntry).filter(
            models.ClassTimetableEntry.course_id.in_(course_ids),
        ).all() if course_ids else []
    else:
        entries = db.query(models.ClassTimetableEntry).filter(
            models.ClassTimetableEntry.lecturer_id == current_user.id,
        ).all()

    # Build iCalendar
    lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Maranatha Risk System//Class Timetable//EN",
        "CALSCALE:GREGORIAN",
        "METHOD:PUBLISH",
    ]

    # Use next Monday as the base date for recurring weekly events
    from datetime import date
    today = date.today()
    days_until_monday = (7 - today.weekday()) % 7
    if days_until_monday == 0:
        days_until_monday = 7
    next_monday = today + timedelta(days=days_until_monday)

    DAY_MAP = {"Monday": 0, "Tuesday": 1, "Wednesday": 2, "Thursday": 3, "Friday": 4, "Saturday": 5, "Sunday": 6}

    for entry in entries:
        day_offset = DAY_MAP.get(entry.day_of_week, 0)
        event_date = next_monday + timedelta(days=day_offset)

        start_time = entry.start_time
        end_time = entry.end_time
        if not start_time or not end_time:
            continue

        # Format: 20260101T090000
        dtstart = f"{event_date.strftime('%Y%m%d')}T{start_time.replace(':', '')}00"
        dtend = f"{event_date.strftime('%Y%m%d')}T{end_time.replace(':', '')}00"

        course = db.query(models.Course).filter(models.Course.id == entry.course_id).first()
        summary = f"{course.course_code} - {course.course_title}" if course else "Class"
        location = entry.venue or ""

        lines.extend([
            "BEGIN:VEVENT",
            f"DTSTART:{dtstart}",
            f"DTEND:{dtend}",
            f"RRULE:FREQ=WEEKLY;COUNT=16",
            f"SUMMARY:{summary}",
            f"LOCATION:{location}",
            f"DESCRIPTION:{entry.entry_type or 'lecture'}",
            f"UID:{entry.id}@maranatha-risk-system",
            "END:VEVENT",
        ])

    lines.append("END:VCALENDAR")

    ics_content = "\r\n".join(lines)
    return Response(
        content=ics_content,
        media_type="text/calendar",
        headers={"Content-Disposition": "attachment; filename=timetable.ics"},
    )
